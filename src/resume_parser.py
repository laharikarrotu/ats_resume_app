"""
Resume Parser: Extract structured data from PDF and DOCX resume files.

This module parses resumes to extract:
- Contact information (name, email, phone, LinkedIn, GitHub, location)
- Education (degrees, universities, GPAs, coursework)
- Technical skills (organized by category)
- Work experience (company, role, dates, bullet points)
- Projects (name, technologies, descriptions)
- Certifications
"""

import re
from pathlib import Path
from typing import List, Optional, Dict, Any

from docx import Document
import pdfplumber
from pypdf import PdfReader

from .models import (
    ResumeData,
    Education,
    Experience,
    Project,
    Certification,
)


def parse_resume(file_path: str) -> ResumeData:
    """
    Parse a resume file (PDF or DOCX) and extract structured data.
    
    Args:
        file_path: Path to the resume file (PDF or DOCX)
    
    Returns:
        ResumeData: Structured resume data
    """
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"Resume file not found: {file_path}")
    
    if path.suffix.lower() == '.pdf':
        text = _extract_text_from_pdf(file_path)
    elif path.suffix.lower() in ['.docx', '.doc']:
        text = _extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {path.suffix}. Only PDF and DOCX are supported.")
    
    # Parse structured data from text
    return _parse_text_to_resume_data(text)


def _extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file using pdfplumber (better for structured text)."""
    text_parts = []
    
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
    except Exception as e:
        # Fallback to pypdf if pdfplumber fails
        reader = PdfReader(file_path)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
    
    return "\n".join(text_parts)


def _extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    doc = Document(file_path)
    text_parts = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
            if row_text:
                text_parts.append(row_text)
    
    return "\n".join(text_parts)


def _parse_text_to_resume_data(text: str) -> ResumeData:
    """
    Parse extracted text into structured ResumeData.
    
    Uses pattern matching and heuristics to identify sections.
    """
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    full_text = '\n'.join(lines)
    
    # Extract contact information
    name = _extract_name(lines)
    email = _extract_email(full_text)
    phone = _extract_phone(full_text)
    linkedin = _extract_linkedin(full_text)
    github = _extract_github(full_text)
    location = _extract_location(full_text)
    
    # Extract sections
    education = _extract_education(full_text, lines)
    skills = _extract_skills(full_text, lines)
    experience = _extract_experience(full_text, lines)
    projects = _extract_projects(full_text, lines)
    certifications = _extract_certifications(full_text, lines)
    
    return ResumeData(
        name=name or "Your Name",
        email=email or "",
        phone=phone or "",
        linkedin=linkedin or "",
        github=github or "",
        location=location or "",
        education=education,
        skills=skills,
        experience=experience,
        projects=projects,
        certifications=certifications,
    )


def _extract_name(lines: List[str]) -> Optional[str]:
    """Extract name (usually first line or first few words)."""
    if not lines:
        return None
    
    # Name is typically the first line, all caps, or title case
    first_line = lines[0]
    
    # If it's all caps or title case and not too long, it's likely the name
    if len(first_line.split()) <= 4 and (first_line.isupper() or first_line.istitle()):
        return first_line
    
    # Check first few lines
    for line in lines[:3]:
        if len(line.split()) <= 4 and line.istitle() and not any(
            keyword in line.lower() for keyword in ['email', 'phone', 'linkedin', 'github', 'resume']
        ):
            return line
    
    return None


def _extract_email(text: str) -> Optional[str]:
    """Extract email address using regex."""
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    match = re.search(email_pattern, text)
    return match.group(0) if match else None


def _extract_phone(text: str) -> Optional[str]:
    """Extract phone number using regex."""
    # Match various phone formats: (123) 456-7890, 123-456-7890, +1 123 456 7890, etc.
    phone_patterns = [
        r'\+?1?\s?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
        r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
    ]
    
    for pattern in phone_patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(0).strip()
    
    return None


def _extract_linkedin(text: str) -> Optional[str]:
    """Extract LinkedIn profile URL."""
    patterns = [
        r'linkedin\.com/in/[\w-]+',
        r'linkedin\.com/[\w-]+',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return f"https://{match.group(0)}"
    
    return None


def _extract_github(text: str) -> Optional[str]:
    """Extract GitHub profile URL."""
    patterns = [
        r'github\.com/[\w-]+',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return f"https://{match.group(0)}"
    
    return None


def _extract_location(text: str) -> Optional[str]:
    """Extract location (city, state or city, country)."""
    # Look for patterns like "City, State" or "City, Country"
    location_pattern = r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*),\s*([A-Z]{2}|[A-Z][a-z]+)'
    match = re.search(location_pattern, text)
    if match:
        return match.group(0)
    
    # Fallback: look for common location keywords
    location_keywords = ['Vacaville', 'CA', 'California', 'Florida', 'New York', 'Texas']
    for keyword in location_keywords:
        if keyword in text:
            # Try to find the full location context
            context_match = re.search(rf'[^,\n]*{re.escape(keyword)}[^,\n]*', text)
            if context_match:
                return context_match.group(0).strip()
    
    return None


def _extract_education(text: str, lines: List[str]) -> List[Education]:
    """Extract education information."""
    education_list = []
    
    # Find education section
    education_start = -1
    for i, line in enumerate(lines):
        if re.search(r'education', line, re.IGNORECASE):
            education_start = i
            break
    
    if education_start == -1:
        return education_list
    
    # Extract education entries (usually 2-3 entries)
    current_education = None
    for i in range(education_start + 1, min(education_start + 20, len(lines))):
        line = lines[i]
        
        # Check if we've hit the next major section
        if re.search(r'(experience|skills|projects|certifications|technical)', line, re.IGNORECASE):
            if current_education:
                education_list.append(current_education)
            break
        
        # Look for degree patterns (Master's, Bachelor's, PhD, etc.)
        degree_match = re.search(
            r'(Master\'?s?|Bachelor\'?s?|PhD|Ph\.D\.|Doctorate|Associate)\s+(?:in|of)?\s+([A-Z][\w\s&]+)',
            line,
            re.IGNORECASE
        )
        
        if degree_match:
            if current_education:
                education_list.append(current_education)
            
            degree = degree_match.group(0)
            current_education = Education(
                degree=degree,
                university="",
                location="",
                dates="",
                gpa="",
                coursework=[]
            )
        elif current_education:
            # Extract university name
            if not current_education.university and len(line.split()) <= 5:
                # Likely university name
                if any(keyword in line.lower() for keyword in ['university', 'college', 'institute', 'tech']):
                    current_education.university = line
            
            # Extract dates (YYYY-YYYY or YYYY - YYYY)
            date_match = re.search(r'(\d{4})\s*[-–]\s*(\d{4}|Present|Current)', line)
            if date_match:
                current_education.dates = date_match.group(0)
            
            # Extract GPA
            gpa_match = re.search(r'GPA[:\s]+([\d.]+)', line, re.IGNORECASE)
            if gpa_match:
                current_education.gpa = gpa_match.group(1)
            
            # Extract coursework
            if 'coursework' in line.lower() or 'courses' in line.lower():
                # Next few lines might be coursework
                coursework_text = ' '.join(lines[i+1:i+3])
                coursework_items = re.findall(r'([A-Z][\w\s&]+)', coursework_text)
                current_education.coursework = coursework_items[:10]  # Limit to 10 items
    
    if current_education:
        education_list.append(current_education)
    
    return education_list


def _extract_skills(text: str, lines: List[str]) -> Dict[str, List[str]]:
    """Extract technical skills organized by category."""
    skills_dict: Dict[str, List[str]] = {}
    
    # Find skills section
    skills_start = -1
    for i, line in enumerate(lines):
        if re.search(r'(technical\s+)?skills?', line, re.IGNORECASE):
            skills_start = i
            break
    
    if skills_start == -1:
        return skills_dict
    
    current_category = "General"
    for i in range(skills_start + 1, min(skills_start + 30, len(lines))):
        line = lines[i]
        
        # Check if we've hit the next major section
        if re.search(r'(experience|projects|certifications|education)', line, re.IGNORECASE):
            break
        
        # Check if this is a category header (bold, all caps, or followed by colon)
        if re.match(r'^[A-Z][\w\s&]+:', line) or (line.isupper() and len(line.split()) <= 3):
            category = line.rstrip(':').strip()
            current_category = category
            skills_dict[current_category] = []
        else:
            # Extract skills from this line (comma or pipe separated)
            if current_category not in skills_dict:
                skills_dict[current_category] = []
            
            # Split by comma, semicolon, or pipe
            skill_items = re.split(r'[,;|]', line)
            for item in skill_items:
                skill = item.strip()
                if skill and len(skill) > 1:
                    skills_dict[current_category].append(skill)
    
    return skills_dict


def _extract_experience(text: str, lines: List[str]) -> List[Experience]:
    """Extract work experience."""
    experience_list = []
    
    # Find experience section
    experience_start = -1
    for i, line in enumerate(lines):
        if re.search(r'(work\s+)?experience|employment|professional\s+experience', line, re.IGNORECASE):
            experience_start = i
            break
    
    if experience_start == -1:
        return experience_list
    
    current_experience = None
    bullets = []
    
    for i in range(experience_start + 1, min(experience_start + 50, len(lines))):
        line = lines[i]
        
        # Check if we've hit the next major section
        if re.search(r'(projects|certifications|education|skills)', line, re.IGNORECASE):
            if current_experience:
                current_experience.bullets = bullets
                experience_list.append(current_experience)
            break
        
        # Look for job title patterns (usually title at company or title | company)
        title_company_match = re.search(
            r'^([A-Z][\w\s&/]+)\s+(?:at|@|\|)\s+([A-Z][\w\s&]+)',
            line
        )
        
        if title_company_match:
            if current_experience:
                current_experience.bullets = bullets
                experience_list.append(current_experience)
            
            title = title_company_match.group(1).strip()
            company = title_company_match.group(2).strip()
            current_experience = Experience(
                title=title,
                company=company,
                dates="",
                bullets=[]
            )
            bullets = []
        elif current_experience:
            # Extract dates
            date_match = re.search(r'(\w+\s+\d{4})\s*[-–]\s*(\w+\s+\d{4}|Present|Current)', line)
            if date_match:
                current_experience.dates = date_match.group(0)
            
            # Extract bullet points (lines starting with •, -, *, or indented)
            if re.match(r'^[•\-\*]\s+', line) or line.startswith('  '):
                bullet = re.sub(r'^[•\-\*]\s+', '', line).strip()
                if bullet:
                    bullets.append(bullet)
    
    if current_experience:
        current_experience.bullets = bullets
        experience_list.append(current_experience)
    
    return experience_list


def _extract_projects(text: str, lines: List[str]) -> List[Project]:
    """Extract projects."""
    projects_list = []
    
    # Find projects section
    projects_start = -1
    for i, line in enumerate(lines):
        if re.search(r'projects?', line, re.IGNORECASE):
            projects_start = i
            break
    
    if projects_start == -1:
        return projects_list
    
    current_project = None
    
    for i in range(projects_start + 1, min(projects_start + 40, len(lines))):
        line = lines[i]
        
        # Check if we've hit the next major section
        if re.search(r'(certifications|education|skills|experience)', line, re.IGNORECASE):
            if current_project:
                projects_list.append(current_project)
            break
        
        # Look for project name (usually bold or title case, may have parentheses)
        project_match = re.search(r'^([A-Z][\w\s&/]+(?:\([^)]+\))?)', line)
        
        if project_match and len(line.split()) <= 8:
            if current_project:
                projects_list.append(current_project)
            
            project_name = project_match.group(1).strip()
            current_project = Project(
                name=project_name,
                description="",
                technologies=[],
                category=""
            )
        elif current_project:
            # Extract technologies (look for tech keywords or parentheses)
            tech_match = re.search(r'\(([^)]+)\)', line)
            if tech_match:
                techs = [t.strip() for t in tech_match.group(1).split(',')]
                current_project.technologies.extend(techs)
            
            # Extract description (bullet points or paragraphs)
            if re.match(r'^[•\-\*]\s+', line) or (len(line) > 20 and not line.isupper()):
                desc = re.sub(r'^[•\-\*]\s+', '', line).strip()
                if desc:
                    if current_project.description:
                        current_project.description += " " + desc
                    else:
                        current_project.description = desc
    
    if current_project:
        projects_list.append(current_project)
    
    return projects_list


def _extract_certifications(text: str, lines: List[str]) -> List[Certification]:
    """Extract certifications."""
    certifications_list = []
    
    # Find certifications section
    cert_start = -1
    for i, line in enumerate(lines):
        if re.search(r'certifications?', line, re.IGNORECASE):
            cert_start = i
            break
    
    if cert_start == -1:
        return certifications_list
    
    for i in range(cert_start + 1, min(cert_start + 20, len(lines))):
        line = lines[i]
        
        # Check if we've hit the next major section
        if re.search(r'(education|skills|experience|projects)', line, re.IGNORECASE):
            break
        
        # Look for certification patterns
        cert_match = re.search(r'([A-Z][\w\s&/]+(?:Certified|Certification|Associate|Professional))', line)
        if cert_match:
            cert_name = cert_match.group(1).strip()
            
            # Extract year if present
            year_match = re.search(r'\((\d{4})\)', line)
            year = year_match.group(1) if year_match else ""
            
            certifications_list.append(Certification(
                name=cert_name,
                issuer="",
                year=year
            ))
    
    return certifications_list

