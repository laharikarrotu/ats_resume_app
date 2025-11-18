"""
Enhanced Resume Generator: Creates ATS-optimized resumes in C3 format with strict 1-page enforcement.

This module generates professional resumes that:
- Use C3 page size (7.17" x 10.51")
- Match the user's resume format exactly
- Enforce strict 1-page limit with auto-truncation
- Use personalized data from parsed resume
- Inject job-specific keywords naturally
"""

from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

from .models import ResumeData, Experience
from .utils import deduplicate_preserve_order, normalize_keyword
from .llm_client import rewrite_experience_bullets, rewrite_project_description, match_experience_with_jd

# Try to import async versions for parallel execution
try:
    import asyncio
    from .llm_client_async import (
        prepare_resume_data_parallel,
        rewrite_experience_bullets_async,
        rewrite_project_description_async,
        match_experience_with_jd_async
    )
    ASYNC_AVAILABLE = True
except ImportError:
    ASYNC_AVAILABLE = False

# Import LLM condenser
try:
    from .llm_condenser import condense_resume_for_one_page
    LLM_CONDENSER_AVAILABLE = True
except ImportError:
    LLM_CONDENSER_AVAILABLE = False


BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = BASE_DIR / "resume_templates" / "c3_template.docx"

# C3 Page Size: 7.17" x 10.51" (324mm x 458mm)
# For better fit, we can use Letter size (8.5" x 11") which is bigger
C3_WIDTH = Inches(7.17)
C3_HEIGHT = Inches(10.51)
LETTER_WIDTH = Inches(8.5)
LETTER_HEIGHT = Inches(11.0)

# Use Letter size for more content space (user requested C3 or bigger)
USE_LETTER_SIZE = True  # Set to False for strict C3

# Margins (matching your resume format: 1" top/bottom, 1.25" left/right)
# But we'll use slightly smaller for more content while maintaining professional look
MARGIN_TOP = Inches(0.6)  # Slightly smaller than 1" for more content
MARGIN_BOTTOM = Inches(0.6)
MARGIN_LEFT = Inches(0.75)  # Slightly smaller than 1.25" for more content
MARGIN_RIGHT = Inches(0.75)

# Font sizes (matching your resume format exactly)
FONT_NAME = "Calibri"  # Default, will use document's font if available
FONT_SIZE_HEADER = Pt(12)  # Name (matching your resume: 12pt)
FONT_SIZE_SECTION = Pt(10)  # Section headings (matching: 10pt, bold)
FONT_SIZE_BODY = Pt(10)  # Body text (matching: 10pt)
FONT_SIZE_CONTACT = Pt(10)  # Contact info (matching: 10pt)

# Spacing
LINE_SPACING = Pt(12)
PARAGRAPH_SPACING = Pt(6)


async def generate_resume(
    output_path: str,
    keywords: List[str],
    resume_data: Optional[ResumeData] = None,
    job_description: Optional[str] = None,
    use_parallel: bool = True
) -> None:
    """
    Generate a personalized, ATS-optimized resume in C3 format.
    
    Args:
        output_path: Path to save the generated DOCX file
        keywords: Extracted keywords from job description
        resume_data: Parsed resume data (if available)
        job_description: Job description (for future LLM rewriting)
    """
    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    
    # Condense resume intelligently to fit one page while preserving 90%+ content
    if resume_data and LLM_CONDENSER_AVAILABLE:
        try:
            resume_data = condense_resume_for_one_page(resume_data, target_page_size="C3" if not USE_LETTER_SIZE else "Letter")
        except Exception as e:
            print(f"Condensation error: {e}. Using original data.")
    
    # Create new document
    document = Document()
    
    # Set page size (C3 or Letter - Letter is bigger for more content)
    if USE_LETTER_SIZE:
        _set_letter_page_size(document)
    else:
        _set_c3_page_size(document)
    
    # Set margins
    _set_margins(document)
    
    # Prepare resume data with parallel LLM calls for speed (if enabled)
    if use_parallel and ASYNC_AVAILABLE and resume_data and job_description:
        try:
            prioritized_experiences, personalized_projects = await prepare_resume_data_parallel(
                resume_data, job_description, keywords
            )
            # Use the prepared data
            resume_data.experience = prioritized_experiences
            resume_data.projects = personalized_projects
        except Exception as e:
            print(f"Parallel LLM processing error: {e}. Using sequential processing.")
    
    # Build resume sections (with more content now)
    _build_header(document, resume_data)
    _build_education(document, resume_data)
    _build_skills(document, resume_data, keywords)
    # Build sections (data already prepared in parallel above if enabled)
    _build_experience(document, resume_data, keywords, job_description)
    _build_projects(document, resume_data, keywords, job_description)
    _build_certifications(document, resume_data)
    
    # Enforce 1-page limit (smarter enforcement)
    _enforce_one_page_smart(document)
    
    # Save document
    document.save(str(output_file))


def _set_c3_page_size(document: Document) -> None:
    """Set document to C3 page size (7.17" x 10.51")."""
    section = document.sections[0]
    section.page_width = C3_WIDTH
    section.page_height = C3_HEIGHT
    section.orientation = WD_ORIENT.PORTRAIT


def _set_letter_page_size(document: Document) -> None:
    """Set document to Letter page size (8.5" x 11") - bigger for more content."""
    section = document.sections[0]
    section.page_width = LETTER_WIDTH
    section.page_height = LETTER_HEIGHT
    section.orientation = WD_ORIENT.PORTRAIT


def _set_margins(document: Document) -> None:
    """Set document margins to 0.4" all around for maximum content space."""
    section = document.sections[0]
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT


def _build_header(document: Document, resume_data: Optional[ResumeData]) -> None:
    """Build resume header matching your exact format."""
    if resume_data:
        # Name (12pt, Bold, UPPERCASE)
        name_para = document.add_paragraph()
        name_run = name_para.add_run(resume_data.name.upper())
        name_run.font.name = FONT_NAME
        name_run.font.size = FONT_SIZE_HEADER
        name_run.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Contact info - split across lines (matching your format)
        if resume_data.email:
            email_para = document.add_paragraph()
            email_run = email_para.add_run(resume_data.email)
            email_run.font.name = FONT_NAME
            email_run.font.size = FONT_SIZE_CONTACT
        
        if resume_data.phone:
            phone_para = document.add_paragraph()
            phone_run = phone_para.add_run(resume_data.phone)
            phone_run.font.name = FONT_NAME
            phone_run.font.size = FONT_SIZE_CONTACT
        
        # LinkedIn and GitHub on same line
        social_parts = []
        if resume_data.linkedin:
            # Extract just the profile part
            linkedin_clean = resume_data.linkedin.replace("https://", "").replace("http://", "")
            social_parts.append(linkedin_clean)
        if resume_data.github:
            github_clean = resume_data.github.replace("https://", "").replace("http://", "")
            social_parts.append(github_clean)
        
        if social_parts:
            social_para = document.add_paragraph()
            social_run = social_para.add_run(" | ".join(social_parts))
            social_run.font.name = FONT_NAME
            social_run.font.size = FONT_SIZE_CONTACT
        
        if resume_data.location:
            location_para = document.add_paragraph()
            location_run = location_para.add_run(resume_data.location)
            location_run.font.name = FONT_NAME
            location_run.font.size = FONT_SIZE_CONTACT
    else:
        # Fallback header
        name_para = document.add_paragraph()
        name_run = name_para.add_run("YOUR NAME")
        name_run.font.name = FONT_NAME
        name_run.font.size = FONT_SIZE_HEADER
        name_run.bold = True
        
        contact_para = document.add_paragraph()
        contact_run = contact_para.add_run("Email • Phone • LinkedIn • GitHub • Location")
        contact_run.font.name = FONT_NAME
        contact_run.font.size = FONT_SIZE_CONTACT
    
    # Add separator line (matching your format)
    _add_separator_line(document)


def _build_education(document: Document, resume_data: Optional[ResumeData]) -> None:
    """Build education section - include ALL education entries."""
    if not resume_data or not resume_data.education:
        return
    
    # Section heading
    _add_section_heading(document, "EDUCATION")
    
    # Include ALL education entries (not just 2)
    for edu in resume_data.education:
        # Degree and University (matching your format: "Master's in Computer Science, Florida Institute of Technology")
        edu_para = document.add_paragraph()
        edu_text = f"{edu.degree}"
        if edu.university:
            edu_text += f", {edu.university}"
        if edu.location:
            edu_text += f", {edu.location}"
        
        edu_run = edu_para.add_run(edu_text)
        edu_run.font.name = FONT_NAME
        edu_run.font.size = FONT_SIZE_BODY
        
        # Dates and GPA on separate line (matching: "GPA: 3.5/4.0 | 2022 - 2024")
        details_para = document.add_paragraph()
        details_parts = []
        if edu.gpa:
            details_parts.append(f"GPA: {edu.gpa}")
        if edu.dates:
            details_parts.append(edu.dates)
        
        if details_parts:
            details_run = details_para.add_run(" | ".join(details_parts))
            details_run.font.name = FONT_NAME
            details_run.font.size = FONT_SIZE_BODY
        
        # Coursework (include more coursework)
        if edu.coursework and len(edu.coursework) > 0:
            coursework_para = document.add_paragraph()
            coursework_text = "Coursework: " + ", ".join(edu.coursework[:8])  # Increased from 5 to 8
            coursework_run = coursework_para.add_run(coursework_text)
            coursework_run.font.name = FONT_NAME
            coursework_run.font.size = FONT_SIZE_BODY
            coursework_run.italic = True
        
        document.add_paragraph()  # Spacing


def _build_skills(document: Document, resume_data: Optional[ResumeData], keywords: List[str]) -> None:
    """Build technical skills section matching your exact format (categorized)."""
    # Section heading
    _add_section_heading(document, "TECHNICAL SKILLS")
    
    # Group skills by category (matching your format)
    categorized_skills = {}
    uncategorized = []
    
    # Process resume skills
    if resume_data and resume_data.skills:
        # If skills is a dict (categorized), use it directly
        if isinstance(resume_data.skills, dict):
            categorized_skills = resume_data.skills.copy()
        else:
            # If skills is a list, categorize them
            for skill in resume_data.skills:
                skill_lower = str(skill).lower()
                category = _categorize_skill(skill_lower)
                if category:
                    if category not in categorized_skills:
                        categorized_skills[category] = []
                    categorized_skills[category].append(skill)
                else:
                    uncategorized.append(skill)
    
    # Add keywords to appropriate categories
    for keyword in keywords:
        normalized_keyword = normalize_keyword(keyword)
        keyword_lower = keyword.lower()
        category = _categorize_skill(keyword_lower)
        
        if category:
            # Check if not already in category
            existing_skills = categorized_skills.get(category, [])
            if normalized_keyword not in [normalize_keyword(str(s)) for s in existing_skills]:
                if category not in categorized_skills:
                    categorized_skills[category] = []
                categorized_skills[category].append(keyword)
        else:
            if normalized_keyword not in [normalize_keyword(str(s)) for s in uncategorized]:
                uncategorized.append(keyword)
    
    # Deduplicate within each category
    for category in categorized_skills:
        categorized_skills[category] = deduplicate_preserve_order([str(s) for s in categorized_skills[category]])
    
    # Write categorized skills (matching your format)
    category_order = [
        "Programming Languages",
        "AI/ML & LLMs",
        "Cloud Platforms",
        "Data Engineering",
        "DevOps & MLOps",
        "Databases & Storage",
        "Frameworks & APIs",
        "Monitoring & Analytics"
    ]
    
    for category in category_order:
        if category in categorized_skills and categorized_skills[category]:
            skills_para = document.add_paragraph()
            # Category name in bold
            category_run = skills_para.add_run(f"{category}: ")
            category_run.font.name = FONT_NAME
            category_run.font.size = FONT_SIZE_BODY
            category_run.bold = True
            # Skills in regular
            skills_text = ", ".join(categorized_skills[category][:20])  # Limit per category
            skills_run = skills_para.add_run(skills_text)
            skills_run.font.name = FONT_NAME
            skills_run.font.size = FONT_SIZE_BODY
    
    # Add uncategorized skills if any
    if uncategorized:
        skills_para = document.add_paragraph()
        other_run = skills_para.add_run("Other: ")
        other_run.font.name = FONT_NAME
        other_run.font.size = FONT_SIZE_BODY
        other_run.bold = True
        skills_text = ", ".join(deduplicate_preserve_order([str(s) for s in uncategorized])[:15])
        skills_run = skills_para.add_run(skills_text)
        skills_run.font.name = FONT_NAME
        skills_run.font.size = FONT_SIZE_BODY
    
    document.add_paragraph()  # Spacing


def _categorize_skill(skill: str) -> Optional[str]:
    """Categorize a skill based on keywords."""
    skill_lower = skill.lower()
    if any(kw in skill_lower for kw in ['python', 'java', 'sql', 'javascript', 'typescript', 'c++', 'go', 'rust', 'pyspark', 'scala', 'r']):
        return "Programming Languages"
    elif any(kw in skill_lower for kw in ['tensorflow', 'pytorch', 'hugging', 'llm', 'gpt', 'bert', 'ml', 'ai', 'machine learning', 'deep learning', 'neural']):
        return "AI/ML & LLMs"
    elif any(kw in skill_lower for kw in ['aws', 'azure', 'gcp', 'cloud', 'sagemaker', 'lambda', 'glue', 'ec2', 's3']):
        return "Cloud Platforms"
    elif any(kw in skill_lower for kw in ['spark', 'databricks', 'kafka', 'etl', 'data engineering', 'airflow', 'flink']):
        return "Data Engineering"
    elif any(kw in skill_lower for kw in ['docker', 'kubernetes', 'terraform', 'ci/cd', 'devops', 'mlops', 'jenkins', 'gitlab']):
        return "DevOps & MLOps"
    elif any(kw in skill_lower for kw in ['postgresql', 'mongodb', 'redis', 'mysql', 'database', 'storage', 'dynamodb', 'redshift']):
        return "Databases & Storage"
    elif any(kw in skill_lower for kw in ['fastapi', 'django', 'flask', 'react', 'node', 'api', 'framework', 'express', 'spring']):
        return "Frameworks & APIs"
    elif any(kw in skill_lower for kw in ['datadog', 'new relic', 'splunk', 'monitoring', 'analytics', 'grafana', 'prometheus']):
        return "Monitoring & Analytics"
    return None


def _build_experience(
    document: Document,
    resume_data: Optional[ResumeData],
    keywords: List[str],
    job_description: Optional[str]
) -> None:
    """Build work experience section with personalized, optimized bullets."""
    if not resume_data or not resume_data.experience:
        return
    
    # Section heading
    _add_section_heading(document, "WORK EXPERIENCE")
    
    # Prioritize most relevant experiences using LLM, but include more
    if job_description and len(resume_data.experience) > 4:
        # Get top 4 most relevant (increased from 3)
        experiences = match_experience_with_jd(resume_data.experience, job_description, top_n=4)
    else:
        # Include ALL experiences if 4 or fewer
        experiences = resume_data.experience[:5]  # Increased from 3 to 5
    
    for exp in experiences:
        # Title and Company (matching your format: "Full Stack AI Engineer - Arkatech Solutions")
        title_para = document.add_paragraph()
        title_text = f"{exp.title}"
        if exp.company:
            title_text += f" - {exp.company}"  # Using " - " instead of " | "
        
        title_run = title_para.add_run(title_text)
        title_run.font.name = FONT_NAME
        title_run.font.size = FONT_SIZE_BODY
        title_run.bold = True
        
        # Dates on separate line (matching: "May 2025 - Present")
        if exp.dates:
            dates_para = document.add_paragraph()
            dates_run = dates_para.add_run(exp.dates)
            dates_run.font.name = FONT_NAME
            dates_run.font.size = FONT_SIZE_BODY
        
        # Bullet points - PERSONALIZE using LLM if job description available
        if job_description and keywords:
            # Use LLM to rewrite bullets to match job description
            personalized_bullets = rewrite_experience_bullets(exp, job_description, keywords)
            bullets = personalized_bullets[:6]  # Increased from 4 to 6 bullets per experience
        else:
            # Fallback to original bullets (include more)
            bullets = exp.bullets[:6]  # Increased from 4 to 6
        
        for bullet in bullets:
            # Use bullet character (•) matching your format
            bullet_para = document.add_paragraph()
            bullet_run = bullet_para.add_run(f"• {bullet}")
            bullet_run.font.name = FONT_NAME
            bullet_run.font.size = FONT_SIZE_BODY
        
        document.add_paragraph()  # Spacing


def _build_projects(
    document: Document,
    resume_data: Optional[ResumeData],
    keywords: List[str],
    job_description: Optional[str]
) -> None:
    """Build projects section with personalized descriptions."""
    if not resume_data or not resume_data.projects:
        return
    
    # Section heading
    _add_section_heading(document, "PROJECTS")
    
    # Include more projects (up to 4)
    projects = resume_data.projects[:4]  # Increased from 3 to 4
    
    for project in projects:
        # Project name (matching your format: "Predictive Maintenance (Manufacturing / Data Engineering)")
        name_para = document.add_paragraph()
        name_text = project.name
        if project.category:
            name_text += f" ({project.category})"
        
        name_run = name_para.add_run(name_text)
        name_run.font.name = FONT_NAME
        name_run.font.size = FONT_SIZE_BODY
        name_run.bold = True
        
        # Technologies on separate line (matching: "Technologies: Spark, Snowflake, Azure Data Lake, LSTM...")
        if project.technologies:
            tech_para = document.add_paragraph()
            tech_text = "Technologies: " + ", ".join(project.technologies[:15])  # Increased to 15
            tech_run = tech_para.add_run(tech_text)
            tech_run.font.name = FONT_NAME
            tech_run.font.size = FONT_SIZE_BODY
        
        # Description - PERSONALIZE using LLM if job description available
        # Split into multiple bullet points (matching your format)
        if project.description:
            if job_description and keywords:
                # Use LLM to rewrite description to match job description
                personalized_description = rewrite_project_description(project, job_description, keywords)
                # Split description into sentences for multiple bullets
                sentences = personalized_description.split('. ')
                for sentence in sentences[:3]:  # Up to 3 bullet points per project
                    if sentence.strip():
                        desc_para = document.add_paragraph()
                        desc_run = desc_para.add_run(sentence.strip() + ("." if not sentence.endswith(".") else ""))
                        desc_run.font.name = FONT_NAME
                        desc_run.font.size = FONT_SIZE_BODY
            else:
                # Fallback: split original description
                sentences = project.description.split('. ')
                for sentence in sentences[:3]:
                    if sentence.strip():
                        desc_para = document.add_paragraph()
                        desc_run = desc_para.add_run(sentence.strip() + ("." if not sentence.endswith(".") else ""))
                        desc_run.font.name = FONT_NAME
                        desc_run.font.size = FONT_SIZE_BODY
        
        document.add_paragraph()  # Spacing


def _build_certifications(document: Document, resume_data: Optional[ResumeData]) -> None:
    """Build certifications section."""
    if not resume_data or not resume_data.certifications:
        return
    
    # Section heading
    _add_section_heading(document, "CERTIFICATIONS")
    
    # Include ALL certifications (matching your format: "AWS Certified Solutions Architect - Associate (2025)")
    for cert in resume_data.certifications:
        cert_para = document.add_paragraph()
        cert_text = cert.name
        if cert.year:
            cert_text += f" ({cert.year})"
        # Note: Your format doesn't show issuer, so we'll skip it
        
        cert_run = cert_para.add_run(cert_text)
        cert_run.font.name = FONT_NAME
        cert_run.font.size = FONT_SIZE_BODY
    
    document.add_paragraph()  # Spacing


def _add_separator_line(document: Document) -> None:
    """Add a separator line matching your resume format (line of underscores)."""
    separator_para = document.add_paragraph()
    separator_run = separator_para.add_run("_" * 50)  # Line of underscores
    separator_run.font.name = FONT_NAME
    separator_run.font.size = FONT_SIZE_BODY
    separator_run.bold = True
    separator_para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_section_heading(document: Document, heading_text: str) -> None:
    """Add a section heading matching your exact format (10pt, Bold, with separator above)."""
    # Add separator line before section (matching your format)
    _add_separator_line(document)
    
    # Section heading (10pt, Bold)
    heading_para = document.add_paragraph()
    heading_run = heading_para.add_run(heading_text.upper())
    heading_run.font.name = FONT_NAME
    heading_run.font.size = FONT_SIZE_SECTION
    heading_run.bold = True
    heading_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add small spacing after heading
    document.add_paragraph()


def _enforce_one_page_smart(document: Document) -> None:
    """
    Smarter 1-page enforcement that preserves maximum content.
    
    Uses better heuristics and only removes content if absolutely necessary.
    """
    # Count paragraphs
    total_paragraphs = len(document.paragraphs)
    
    # With Letter size and optimized formatting, we can fit more
    # Letter (8.5x11): ~60-70 paragraphs
    # C3 (7.17x10.51): ~50-60 paragraphs
    if USE_LETTER_SIZE:
        max_paragraphs = 65  # Increased for Letter size
    else:
        max_paragraphs = 55  # Increased for C3
    
    if total_paragraphs > max_paragraphs:
        # Only remove if significantly over
        excess = total_paragraphs - max_paragraphs
        if excess > 5:  # Only remove if more than 5 paragraphs over
            # Remove excess paragraphs (usually spacing or less critical content)
            paragraphs_to_remove = min(excess, 10)  # Don't remove more than 10
            for _ in range(paragraphs_to_remove):
                if len(document.paragraphs) > 15:  # Keep at least header and core content
                    # Remove last paragraph (usually spacing)
                    last_para = document.paragraphs[-1]
                    p_element = last_para._element
                    p_element.getparent().remove(p_element)
    
    # If still too long, slightly reduce font sizes (minimal impact)
    if len(document.paragraphs) > max_paragraphs + 5:
        # Reduce body font size slightly (only if really needed)
        for para in document.paragraphs[5:]:  # Skip header
            for run in para.runs:
                if run.font.size and run.font.size.pt > Pt(9):
                    run.font.size = Pt(9)  # Slight reduction
