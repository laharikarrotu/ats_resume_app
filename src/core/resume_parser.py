"""
Resume Parser: Extract structured data from PDF and DOCX resume files.

This module parses resumes to extract:
- Contact information (name, email, phone, LinkedIn, GitHub, location)
- Education (degrees, universities, GPAs, coursework)
- Technical skills (organized by category)
- Work experience (company, role, dates, bullet points)
- Projects (name, technologies, descriptions)
- Certifications
"""

import re
from pathlib import Path
from typing import List, Optional, Dict, Any

from docx import Document

from ..models import (
    ResumeData,
    Education,
    Experience,
    Project,
    Certification,
)
from ..logger import logger

# Try to import LLM parser, fallback if not available
try:
    from ..llm.parser import parse_resume_with_llm
    LLM_PARSER_AVAILABLE = True
except ImportError:
    LLM_PARSER_AVAILABLE = False

# Smart PDF extractor with space recovery + link extraction
from .pdf_extractor import extract_text_from_pdf as _smart_pdf_extract, extract_pdf_links


def parse_resume(file_path: str) -> ResumeData:
    """
    Parse a resume file (PDF, DOCX, or TXT) and extract structured data.
    
    Args:
        file_path: Path to the resume file (PDF, DOCX, or TXT)
    
    Returns:
        ResumeData: Structured resume data
    """
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"Resume file not found: {file_path}")
    
    pdf_links: dict = {}

    if path.suffix.lower() == '.pdf':
        text = _smart_pdf_extract(file_path)
        pdf_links = extract_pdf_links(file_path)
    elif path.suffix.lower() in ['.docx', '.doc']:
        text = _extract_text_from_docx(file_path)
    elif path.suffix.lower() in ['.txt', '.text']:
        text = _extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {path.suffix}. Supported: PDF, DOCX, TXT.")
    
    logger.info("Extracted %d chars from %s", len(text), path.name)
    result = _parse_text(text)

    # Enrich with PDF hyperlinks if the text only had labels (e.g., "LinkedIn", "GitHub")
    if pdf_links:
        if pdf_links.get("linkedin") and (not result.linkedin or result.linkedin == "LinkedIn"):
            result.linkedin = pdf_links["linkedin"]
        if pdf_links.get("github") and (not result.github or result.github == "GitHub"):
            result.github = pdf_links["github"]

    return result


def parse_resume_from_text(text: str) -> ResumeData:
    """
    Parse resume from raw text (pasted / typed by user).
    
    Args:
        text: Raw resume text content
    
    Returns:
        ResumeData: Structured resume data
    """
    if not text or not text.strip():
        raise ValueError("Resume text is empty.")
    return _parse_text(text.strip())


def _parse_text(text: str) -> ResumeData:
    """
    Smart hybrid parser — regex first (free + instant), LLM only when needed.

    Strategy:
      1. Always run the regex parser first (zero cost, <50ms)
      2. If result looks complete (has name + experience + skills), return it
      3. If result is sparse/incomplete, escalate to LLM for better extraction
    This saves LLM quota for the rewriting step where it matters most.
    """
    # Step 1: Try regex parser first (always free)
    regex_result = _parse_text_to_resume_data(text)

    # Step 2: Check quality — if regex did a good job, skip LLM
    has_name = regex_result.name and regex_result.name != "Your Name"
    has_experience = len(regex_result.experience) > 0
    has_skills = any(len(v) > 0 for v in regex_result.skills.values())
    has_education = len(regex_result.education) > 0
    has_bullets = any(len(exp.bullets) > 0 for exp in regex_result.experience) if has_experience else False

    quality_score = sum([has_name, has_experience, has_skills, has_education])

    # Must have experience with bullets to skip LLM — experience is the most important section
    if quality_score >= 3 and has_experience and has_bullets:
        # Regex did well — no need to spend LLM quota
        logger.info(
            "Regex parser sufficient (quality=%d/4): %s, %d exp, %d projects, %d skills",
            quality_score, regex_result.name, len(regex_result.experience),
            len(regex_result.projects), sum(len(v) for v in regex_result.skills.values()),
        )
        return regex_result

    # Step 3: Regex result is sparse — use LLM for better extraction
    if LLM_PARSER_AVAILABLE:
        try:
            llm_result = parse_resume_with_llm(text)
            logger.info(
                "LLM parser used (regex quality=%d/4): %s, %d exp, %d projects",
                quality_score, llm_result.name, len(llm_result.experience),
                len(llm_result.projects),
            )
            return llm_result
        except Exception as e:
            logger.warning("LLM parser failed: %s — using regex result", e)
            return regex_result
    else:
        logger.info("LLM parser not available, using regex result (quality=%d/4)", quality_score)
        return regex_result


def _extract_text_from_txt(file_path: str) -> str:
    """Extract text from a plain-text (.txt) file."""
    path = Path(file_path)
    return path.read_text(encoding="utf-8", errors="replace")


def _extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    doc = Document(file_path)
    text_parts = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
            if row_text:
                text_parts.append(row_text)
    
    return "\n".join(text_parts)


def _find_section_boundaries(lines: List[str]) -> Dict[str, int]:
    """
    Find the start line index of each major resume section.
    Returns a dict like {'experience': 5, 'education': 30, 'skills': 35, ...}
    """
    section_patterns = {
        "experience": r"^(?:work\s+)?experience|^employment|^professional\s+experience|^relevant\s+experience",
        "education": r"^education",
        "skills": r"^(?:technical\s+)?skills?$",
        "projects": r"^(?:personal\s+|academic\s+|side\s+|key\s+|selected\s+|relevant\s+|notable\s+)?projects?(?:\s*(?:&|and)\s*\w+)?$",
        "certifications": r"^certifications?|^licenses?\s*(?:&|and)?\s*certifications?",
        "summary": r"^(?:professional\s+)?summary|^objective|^profile",
    }
    boundaries: Dict[str, int] = {}
    for i, line in enumerate(lines):
        stripped = line.strip()
        for section, pattern in section_patterns.items():
            if section not in boundaries and re.search(pattern, stripped, re.IGNORECASE):
                boundaries[section] = i
    return boundaries


def _get_section_lines(lines: List[str], boundaries: Dict[str, int], section: str) -> List[str]:
    """Get all lines belonging to a section (from its header to the next section header)."""
    if section not in boundaries:
        return []
    start = max(0, boundaries[section] + 1)  # skip the header/boundary line; clamp to 0
    # Find the next section that starts after this one
    all_starts = sorted(boundaries.values())
    idx = all_starts.index(boundaries[section])
    end = all_starts[idx + 1] if idx + 1 < len(all_starts) else len(lines)
    return lines[start:end]


def _classify_line(line: str) -> Optional[str]:
    """
    Classify a single line by its content pattern — independent of section headers.

    Returns one of: 'contact', 'education', 'experience_header', 'bullet',
    'skill', 'project_header', 'certification', 'section_header', or None.
    """
    stripped = line.strip()
    if not stripped:
        return None

    # ── Section headers (explicit labels like "EDUCATION", "WORK EXPERIENCE") ──
    section_kw = re.match(
        r'^(?:work\s+)?experience$|^education$|^(?:technical\s+)?skills?$|'
        r'^(?:personal\s+|academic\s+|key\s+|selected\s+)?projects?(?:\s*(?:&|and)\s*\w+)?$|'
        r'^certifications?$|^(?:professional\s+)?summary$|^objective$',
        stripped, re.IGNORECASE,
    )
    if section_kw:
        return "section_header"

    # ── Contact line (email, phone, linkedin, url combos) ──
    email_pat = re.search(r'[\w.-]+@[\w.-]+\.\w+', stripped)
    phone_pat = re.search(r'\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}', stripped)
    if email_pat or phone_pat:
        return "contact"

    # ── Bullet point ──
    if re.match(r'^[•\-\*◦▪►]\s+', stripped):
        return "bullet"

    # ── Skill line ("Category: item1, item2, item3") ──
    if re.match(r'^[\w\s/&]+:\s*\w+.*,\s*\w+', stripped):
        return "skill"

    # ── Education-like ("degree ... university" or "university ... degree") ──
    edu_keywords = r'(?:bachelor|master|b\.?s\.?|m\.?s\.?|ph\.?d|associate|diploma|degree|university|college|institute)'
    if re.search(edu_keywords, stripped, re.IGNORECASE):
        return "education"

    # ── Certification-like ("AWS Certified", "Google Cloud", etc.) ──
    cert_keywords = r'(?:certified|certification|certificate|license|credential|aws\s+solutions|google\s+cloud|azure|pmp|comptia|cisco)'
    if re.search(cert_keywords, stripped, re.IGNORECASE) and len(stripped.split()) <= 15:
        return "certification"

    # ── Experience header (Title, Company, Date pattern) ──
    date_pat = r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\.?\s+\d{4}|Present|Current|\d{4}\s*[-–—]\s*(?:\d{4}|Present)'
    if re.search(date_pat, stripped, re.IGNORECASE) and len(stripped.split()) <= 20:
        # Could be experience or education — if it has a company-like pattern, it's experience
        if re.search(r'[|,]\s*[A-Z]', stripped):
            return "experience_header"
        return "experience_header"  # default to experience for date lines

    # ── Project header (Name | Tech or Name — description) ──
    if re.match(r'^[A-Z][\w\s&\'-]+\s*\|\s*', stripped):
        return "project_header"
    if re.match(r'^[A-Z][\w\s&\'-]+\s*[—–]\s*', stripped) and len(stripped.split()) <= 15:
        return "project_header"

    return None


def _smart_section_boundaries(lines: List[str], boundaries: Dict[str, int]) -> Dict[str, int]:
    """
    Enhance section boundaries by analyzing line content patterns.

    If a section wasn't detected by header matching, scan through unassigned
    lines looking for content that reveals the section type.

    Boundary convention: boundary points to the line BEFORE the first content
    line, so _get_section_lines(start = boundary + 1) includes the content.
    When a real header exists, the header IS the boundary line.
    When no header exists, we use (first_content_line - 1) as a virtual boundary.
    """
    missing_sections = [
        s for s in ["experience", "education", "skills", "projects", "certifications"]
        if s not in boundaries
    ]
    if not missing_sections:
        return boundaries

    assigned_starts = set(boundaries.values())

    for i, line in enumerate(lines):
        if i in assigned_starts:
            continue

        classification = _classify_line(line)
        if not classification or classification in ("section_header", "contact", "bullet"):
            continue

        # Determine the section this line belongs to
        section = None
        if classification == "experience_header" and "experience" not in boundaries:
            section = "experience"
        elif classification == "education" and "education" not in boundaries:
            section = "education"
        elif classification == "skill" and "skills" not in boundaries:
            section = "skills"
        elif classification == "project_header" and "projects" not in boundaries:
            section = "projects"
        elif classification == "certification" and "certifications" not in boundaries:
            section = "certifications"

        if section:
            # If previous line is a section header label, use it as the boundary
            if i > 0 and _classify_line(lines[i - 1]) == "section_header":
                boundaries[section] = i - 1
            else:
                # No header — set boundary to (i - 1) so _get_section_lines includes line i
                # Use -1 if i == 0; _get_section_lines handles max(0, ...) via start calc
                boundaries[section] = i - 1
            assigned_starts.add(boundaries[section])

    return boundaries


def _parse_text_to_resume_data(text: str) -> ResumeData:
    """
    Parse extracted text into structured ResumeData using robust regex patterns.
    This is the fallback parser when LLM is unavailable.

    Uses a two-pass approach:
      1. Header-based section detection (fast, exact)
      2. Content-based line classification (smart, fills gaps)
    """
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    full_text = '\n'.join(lines)

    # Pass 1: Find section boundaries by headers
    boundaries = _find_section_boundaries(lines)

    # Pass 2: Smart fill — if sections are missing, classify lines by content
    boundaries = _smart_section_boundaries(lines, boundaries)

    # Extract contact information
    name = _extract_name(lines, boundaries)
    email = _extract_email(full_text)
    phone = _extract_phone(full_text)
    linkedin = _extract_linkedin(full_text)
    github = _extract_github(full_text)
    location = _extract_location(full_text)

    # Extract sections using enhanced boundaries
    education = _extract_education(lines, boundaries)
    skills = _extract_skills(lines, boundaries)
    experience = _extract_experience(lines, boundaries)
    projects = _extract_projects(lines, boundaries)
    certifications = _extract_certifications(lines, boundaries)

    return ResumeData(
        name=name or "Your Name",
        email=email or "",
        phone=phone or "",
        linkedin=linkedin or "",
        github=github or "",
        location=location or "",
        education=education,
        skills=skills,
        experience=experience,
        projects=projects,
        certifications=certifications,
    )


# ═══════════════════════════════════════════════════════════
# Contact Information Extraction
# ═══════════════════════════════════════════════════════════

def _extract_name(lines: List[str], boundaries: Dict[str, int]) -> Optional[str]:
    """Extract name — typically the first non-empty line before any section header."""
    if not lines:
        return None

    # The name is almost always line 0 or within the first 3 lines
    first_section = min(boundaries.values()) if boundaries else len(lines)

    for line in lines[:min(first_section, 5)]:
        # Skip lines that look like contact info
        if "@" in line or "linkedin" in line.lower() or "github" in line.lower():
            continue
        if re.search(r'\d{3}.*\d{4}', line):  # phone number
            continue
        # Name: 2-4 words, no special characters except hyphens
        words = line.split()
        if 1 <= len(words) <= 5:
            # Remove pipe-separated contact info from the line
            clean = line.split("|")[0].strip()
            if clean and not any(kw in clean.lower() for kw in ["email", "phone", "resume", "http"]):
                return clean
    return None


def _extract_email(text: str) -> Optional[str]:
    """Extract email address using regex."""
    match = re.search(r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}', text)
    return match.group(0) if match else None


def _extract_phone(text: str) -> Optional[str]:
    """Extract phone number — handles many formats."""
    patterns = [
        r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
        r'\+?1[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(0).strip()
    return None


def _extract_linkedin(text: str) -> Optional[str]:
    """Extract LinkedIn profile URL."""
    match = re.search(r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+', text, re.IGNORECASE)
    if match:
        url = match.group(0)
        return url if url.startswith("http") else f"https://{url}"
    # Also check for just the word "LinkedIn" as a hyperlink label
    return None


def _extract_github(text: str) -> Optional[str]:
    """Extract GitHub profile URL."""
    match = re.search(r'(?:https?://)?(?:www\.)?github\.com/[\w-]+', text, re.IGNORECASE)
    if match:
        url = match.group(0)
        return url if url.startswith("http") else f"https://{url}"
    return None


def _extract_location(text: str) -> Optional[str]:
    """Extract location (City, State/Country)."""
    # Pattern: "City, ST" or "City, State"
    match = re.search(r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*),\s*([A-Z]{2})\b', text)
    if match:
        return match.group(0)
    # Broader: "City, State Name"
    match = re.search(r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*),\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)', text)
    if match:
        return match.group(0)
    return None


# ═══════════════════════════════════════════════════════════
# Experience Extraction (most critical section)
# ═══════════════════════════════════════════════════════════

def _extract_experience(lines: List[str], boundaries: Dict[str, int]) -> List[Experience]:
    """Extract work experience with robust pattern matching."""
    section_lines = _get_section_lines(lines, boundaries, "experience")
    if not section_lines:
        return []

    experience_list = []
    current: Optional[Experience] = None
    bullets: List[str] = []
    multi_line_bullet = ""  # accumulator for bullets that wrap across lines

    for line in section_lines:
        # Try to match a job title line in common formats:
        #   "Software Engineer | Oracle  May 2025 – Present"
        #   "Software Engineer, Google, Jul 2022 - Present"
        #   "Software Engineer, Google, Mountain View, CA, Jul 2022 - Present"
        #   "Software Engineer at Oracle"

        # Format 1: "Title [|,] Company [|,spaces] Date"
        title_match = re.match(
            r'^([A-Z][\w\s/&-]+?)\s*[|,]\s*([A-Z][\w\s&.-]+?)(?:\s{2,}|\s*[|,]\s*|\s+)((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[\w\s–\-]+(?:Present|\d{4}))?\s*$',
            line
        )
        if not title_match:
            # Format 2: "Title, Company, Location, Date" (skip location)
            title_match2 = re.match(
                r'^([A-Z][\w\s/&-]+?)\s*,\s*([A-Z][\w\s&.-]+?)\s*,\s*(?:[A-Z][\w\s]+,\s*[A-Z]{2}\s*,?\s*)?((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[\w\s–\-]+(?:Present|\d{4}))\s*$',
                line, re.IGNORECASE
            )
            if title_match2:
                title_match = title_match2
        if not title_match:
            # Format 3: "Title at Company"
            title_match = re.match(
                r'^([A-Z][\w\s/&-]+?)\s+(?:at|@)\s+([A-Z][\w\s&.-]+)',
                line
            )

        if title_match:
            # Save previous experience
            if current:
                if multi_line_bullet:
                    bullets.append(multi_line_bullet.strip())
                    multi_line_bullet = ""
                current.bullets = bullets
                experience_list.append(current)

            title = title_match.group(1).strip()
            company = title_match.group(2).strip()
            dates = title_match.group(3).strip() if title_match.lastindex and title_match.lastindex >= 3 and title_match.group(3) else ""

            current = Experience(title=title, company=company, dates=dates, bullets=[])
            bullets = []
            continue

        # If we have a current experience, look for dates on their own line
        if current and not current.dates:
            date_match = re.search(
                r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\.?\s+\d{4})\s*[-–—]\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\.?\s+\d{4}|Present|Current)',
                line, re.IGNORECASE
            )
            if date_match:
                current.dates = date_match.group(0)
                continue

        # Bullet points
        if current:
            bullet_match = re.match(r'^[•\-\*◦▪►]\s*(.*)', line)
            if bullet_match:
                # Start of a new bullet
                if multi_line_bullet:
                    bullets.append(multi_line_bullet.strip())
                multi_line_bullet = bullet_match.group(1).strip()
            elif multi_line_bullet and line and not line[0].isupper():
                # Continuation of a wrapped bullet (starts lowercase)
                multi_line_bullet += " " + line
            elif multi_line_bullet and line:
                # Could be a continuation of a bullet that wraps to a new line
                # Check if it looks like a continuation (doesn't start a new entry)
                if not re.match(r'^[A-Z][\w\s/&-]+?\s*[|,@]', line):
                    multi_line_bullet += " " + line
                else:
                    bullets.append(multi_line_bullet.strip())
                    multi_line_bullet = ""

    # Flush last experience
    if current:
        if multi_line_bullet:
            bullets.append(multi_line_bullet.strip())
        current.bullets = bullets
        experience_list.append(current)

    return experience_list


# ═══════════════════════════════════════════════════════════
# Education Extraction
# ═══════════════════════════════════════════════════════════

def _extract_education(lines: List[str], boundaries: Dict[str, int]) -> List[Education]:
    """Extract education entries."""
    section_lines = _get_section_lines(lines, boundaries, "education")
    if not section_lines:
        return []

    education_list = []
    current: Optional[Education] = None

    for line in section_lines:
        # Match degree patterns
        # "Master's in Computer Science — Florida Institute of Technology, Melbourne, FL (2024)"
        # "Bachelor of Science in CS, University Name, 2020"
        degree_match = re.search(
            r"(Master'?s?|Bachelor'?s?|PhD|Ph\.D\.?|Doctorate|Associate'?s?|B\.S\.?|M\.S\.?|B\.A\.?|M\.A\.?|MBA)"
            r"(?:\s+(?:of|in|degree\s+in))?\s*([\w\s&/,]+)",
            line, re.IGNORECASE
        )

        if degree_match:
            if current:
                education_list.append(current)

            degree_type = degree_match.group(1)
            rest = degree_match.group(2).strip() if degree_match.group(2) else ""

            # Try to split "Computer Science — University Name, City, State (Year)"
            parts = re.split(r'\s*[—–-]\s*', rest, maxsplit=1)
            major = parts[0].strip().rstrip(",")
            university = ""
            location = ""
            dates = ""

            if len(parts) > 1:
                uni_part = parts[1]
                # Extract year in parens
                year_match = re.search(r'\((\d{4})\)', uni_part)
                if year_match:
                    dates = year_match.group(1)
                    uni_part = uni_part[:year_match.start()].strip().rstrip(",")

                # Extract location (last "City, ST" pattern)
                loc_match = re.search(r',\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*),?\s*([A-Z]{2})?\s*$', uni_part)
                if loc_match:
                    location = loc_match.group(0).strip().lstrip(",").strip()
                    uni_part = uni_part[:loc_match.start()].strip().rstrip(",")

                university = uni_part.strip()

            current = Education(
                degree=f"{degree_type} in {major}" if major else degree_type,
                university=university,
                location=location,
                dates=dates,
                gpa="",
                coursework=[]
            )
            continue

        if current:
            # GPA
            gpa_match = re.search(r'GPA[:\s]+([\d.]+)', line, re.IGNORECASE)
            if gpa_match:
                current.gpa = gpa_match.group(1)

            # Coursework
            if re.search(r'coursework|courses', line, re.IGNORECASE):
                # Extract items after the colon
                coursework_text = re.sub(r'.*(?:coursework|courses)\s*:?\s*', '', line, flags=re.IGNORECASE)
                if coursework_text:
                    items = [c.strip() for c in re.split(r'[,;•]', coursework_text) if c.strip()]
                    current.coursework = items

            # Dates on own line
            if not current.dates:
                date_match = re.search(r'(\d{4})\s*[-–—]\s*(\d{4}|Present)', line)
                if date_match:
                    current.dates = date_match.group(0)

    if current:
        education_list.append(current)

    return education_list


# ═══════════════════════════════════════════════════════════
# Skills Extraction
# ═══════════════════════════════════════════════════════════

def _extract_skills(lines: List[str], boundaries: Dict[str, int]) -> Dict[str, List[str]]:
    """
    Extract skills organized by category.

    Handles formats:
      • "Languages: Python, Java, JavaScript"  (one category per line)
      • "Languages: Python, Java • Cloud: AWS"  (multiple on one line)
      • "Python, Java, JavaScript"              (no category)
    """
    section_lines = _get_section_lines(lines, boundaries, "skills")
    if not section_lines:
        return {}

    skills_dict: Dict[str, List[str]] = {}
    current_category = "General"

    for line in section_lines:
        # Skip bullet-style lines (those belong to experience, not skills)
        if re.match(r'^[•\-\*◦▪►]\s+', line):
            continue

        # Check if line has one or more "Category: items" patterns
        # e.g., "Languages: Python, Java • Cloud: AWS, GCP"
        cat_parts = re.split(r'(?:\s*[•|]\s+|\s{3,})(?=[A-Z][\w/&\s]+?:\s)', line)

        for part in cat_parts:
            cat_match = re.match(r'^([A-Z][\w/&\s]+?):\s*(.*)', part.strip())
            if cat_match:
                current_category = cat_match.group(1).strip()
                raw = cat_match.group(2).strip()
                items = [s.strip() for s in re.split(r'[,;]', raw) if s.strip() and len(s.strip()) > 1]
                # Filter out non-skill content
                items = [s for s in items if not re.match(r'^[•|]\s*$', s)]
                if items:
                    if current_category in skills_dict:
                        skills_dict[current_category].extend(items)
                    else:
                        skills_dict[current_category] = items
            else:
                # No category label — append to current category
                items = [s.strip() for s in re.split(r'[,;|]', part.strip()) if s.strip() and len(s.strip()) > 1]
                if items:
                    if current_category not in skills_dict:
                        skills_dict[current_category] = []
                    skills_dict[current_category].extend(items)

    return skills_dict


# ═══════════════════════════════════════════════════════════
# Projects Extraction
# ═══════════════════════════════════════════════════════════

def _extract_projects(lines: List[str], boundaries: Dict[str, int]) -> List[Project]:
    """
    Extract projects from the Projects section.

    Handles common formats:
      • "Project Name — subtitle (url.com)"
      • "Project Name | Tech1, Tech2, Tech3"
      • "Project Name  (Jan 2024 – Mar 2024)"
      • Followed by bullet points describing the project
    """
    section_lines = _get_section_lines(lines, boundaries, "projects")
    if not section_lines:
        return []

    projects_list: List[Project] = []
    current: Optional[Project] = None
    description_parts: List[str] = []

    def _save_current():
        """Flush the current project into the list."""
        nonlocal current, description_parts
        if current:
            if description_parts:
                current.description = " ".join(description_parts).strip()
            projects_list.append(current)
            current = None
            description_parts = []

    for line in section_lines:
        is_bullet = bool(re.match(r'^[•\-\*◦▪►]', line))

        # ── Try to detect a project header line ──
        # Format 1: "Name — subtitle" or "Name – subtitle" or "Name - subtitle"
        header_dash = re.match(
            r'^([A-Z][\w\s&\'-]+?)(?:\s*[—–|-]\s*(.+))?$', line
        )
        # Format 2: "Name | Tech1, Tech2" (pipe separated)
        header_pipe = re.match(
            r'^([A-Z][\w\s&\'-]+?)\s*\|\s*(.+)$', line
        )
        # Format 3: "Name  (dates)" or just "Name"
        header_plain = re.match(
            r'^([A-Z][\w\s&\'-]+?)(?:\s*\(([^)]+)\))?$', line
        )

        is_header = (
            not is_bullet
            and not line.startswith(" ")
            and len(line.split()) <= 20
            and (header_dash or header_pipe or header_plain)
        )

        if is_header:
            _save_current()

            technologies: List[str] = []
            subtitle = ""

            if header_pipe:
                name = header_pipe.group(1).strip()
                # Pipe part is usually tech list
                tech_str = header_pipe.group(2).strip()
                technologies = [t.strip() for t in re.split(r'[,;]', tech_str) if t.strip()]
            elif header_dash and header_dash.group(2):
                name = header_dash.group(1).strip()
                subtitle = header_dash.group(2).strip()
            elif header_plain:
                name = header_plain.group(1).strip()
                subtitle = header_plain.group(2).strip() if header_plain.group(2) else ""
            else:
                name = line.strip()

            # Extract URL from subtitle/name
            url = ""
            url_match = re.search(r'\(?([a-z][\w.-]+\.[a-z]{2,})\)?', subtitle)
            if url_match:
                url = url_match.group(1)
                subtitle = subtitle[:url_match.start()].strip().rstrip("(").strip()

            # Extract techs from parenthesized list in subtitle
            if not technologies and subtitle:
                tech_paren = re.search(r'(?:using|with|built\s+with|tech(?:nologies)?:?)\s+(.+)', subtitle, re.IGNORECASE)
                if tech_paren:
                    technologies = [t.strip() for t in re.split(r'[,;+]', tech_paren.group(1)) if t.strip() and len(t.strip()) > 1]
                    subtitle = subtitle[:tech_paren.start()].strip()

            current = Project(
                name=name,
                description=subtitle,
                technologies=technologies,
                category=""
            )
        elif current:
            # Bullet or continuation line — belongs to current project
            clean = re.sub(r'^[•\-\*◦▪►]\s*', '', line).strip()
            if clean:
                description_parts.append(clean)
                # Extract technologies mentioned inline
                tech_matches = re.findall(
                    r'(?:using|with|built\s+with|technologies?:?|powered\s+by)\s+([^.]+?)(?:\.|$)',
                    clean, re.IGNORECASE,
                )
                for tm in tech_matches:
                    techs = [t.strip() for t in re.split(r'[,;+]', tm) if t.strip() and len(t.strip()) > 1]
                    for tech in techs:
                        if tech not in current.technologies:
                            current.technologies.append(tech)

    _save_current()
    return projects_list


# ═══════════════════════════════════════════════════════════
# Certifications Extraction
# ═══════════════════════════════════════════════════════════

def _extract_certifications(lines: List[str], boundaries: Dict[str, int]) -> List[Certification]:
    """Extract certifications."""
    section_lines = _get_section_lines(lines, boundaries, "certifications")
    if not section_lines:
        return []

    certifications_list: List[Certification] = []

    for line in section_lines:
        clean = re.sub(r'^[•\-\*◦▪►]\s*', '', line).strip()
        if not clean or len(clean) < 5:
            continue

        # Extract year
        year_match = re.search(r'\(?(\d{4})\)?', clean)
        year = year_match.group(1) if year_match else ""

        # Try to split "Name | Issuer" or "Name – Issuer" or "Name, Issuer"
        parts = re.split(r'\s*[|–—]\s*', clean)
        name = parts[0].strip()
        issuer = parts[1].strip() if len(parts) > 1 else ""

        # Clean up: remove "View Credential" etc.
        name = re.sub(r'\s*\|?\s*View\s+Credential.*', '', name, flags=re.IGNORECASE).strip()
        issuer = re.sub(r'\s*\|?\s*View\s+Credential.*', '', issuer, flags=re.IGNORECASE).strip()

        if name:
            certifications_list.append(Certification(name=name, issuer=issuer, year=year))

    return certifications_list

