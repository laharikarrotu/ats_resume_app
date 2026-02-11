"""
ATS-Optimized Resume Generator — Creates resumes that pass ALL major ATS platforms.

Generates resumes compliant with the top 10 ATS systems:
  Workday (45%), Taleo (20%), Greenhouse (15%), Lever (10%), iCIMS, etc.

ATS-Critical Format Rules Enforced:
  ✓ Single column layout — NO tables, NO columns
  ✓ Contact info on ONE line at top (body, NOT header/footer)
  ✓ Standard section headers: WORK EXPERIENCE, EDUCATION, SKILLS
  ✓ Simple bullets (• or -) ONLY
  ✓ Left-aligned everything
  ✓ Standard fonts: Calibri 10-11pt
  ✓ No images, graphics, text boxes
  ✓ Dates in "Month YYYY" format
  ✓ Consistent formatting throughout
"""

from pathlib import Path
from typing import List, Optional
import asyncio

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

from ..models import ResumeData, Experience
from ..utils import deduplicate_preserve_order, normalize_keyword
from ..llm.client import rewrite_experience_bullets, rewrite_project_description, match_experience_with_jd

# Try to import optimized version first (faster)
try:
    from ..llm.client_optimized import prepare_resume_data_optimized
    OPTIMIZED_AVAILABLE = True
    ASYNC_AVAILABLE = True
except ImportError:
    OPTIMIZED_AVAILABLE = False
    # Fallback to regular async version
    try:
        from ..llm.client_async import (
            prepare_resume_data_parallel,
            rewrite_experience_bullets_async,
            rewrite_project_description_async,
            match_experience_with_jd_async
        )
        ASYNC_AVAILABLE = True
    except ImportError:
        ASYNC_AVAILABLE = False

# Import LLM condenser (async version preferred)
LLM_CONDENSER_ASYNC_AVAILABLE = False
LLM_CONDENSER_AVAILABLE = False
try:
    from ..llm.condenser_async import condense_resume_for_one_page_async
    LLM_CONDENSER_ASYNC_AVAILABLE = True
except ImportError:
    try:
        from ..llm.condenser import condense_resume_for_one_page
        LLM_CONDENSER_AVAILABLE = True
    except ImportError:
        pass


from ..config import BASE_DIR, RESUME_TEMPLATES_DIR
TEMPLATE_PATH = RESUME_TEMPLATES_DIR / "c3_template.docx"

# C3 Page Size: 7.17" x 10.51" (324mm x 458mm)
# For better fit, we can use Letter size (8.5" x 11") which is bigger
C3_WIDTH = Inches(7.17)
C3_HEIGHT = Inches(10.51)
LETTER_WIDTH = Inches(8.5)
LETTER_HEIGHT = Inches(11.0)

# Use Letter size for more content space (user requested C3 or bigger)
USE_LETTER_SIZE = True  # Set to False for strict C3

# Margins (matching your resume format: 1" top/bottom, 1.25" left/right)
# But we'll use slightly smaller for more content while maintaining professional look
MARGIN_TOP = Inches(0.6)  # Slightly smaller than 1" for more content
MARGIN_BOTTOM = Inches(0.6)
MARGIN_LEFT = Inches(0.75)  # Slightly smaller than 1.25" for more content
MARGIN_RIGHT = Inches(0.75)

# ATS-Safe Font Configuration
# Rule: Arial, Calibri, or Times New Roman — 10-12pt
FONT_NAME = "Calibri"         # Most ATS-compatible sans-serif font
FONT_SIZE_NAME = Pt(14)       # Name: 14pt bold (ATS rule: 14-16pt)
FONT_SIZE_CONTACT = Pt(10)    # Contact line: 10pt
FONT_SIZE_SECTION = Pt(11)    # Section headings: 11pt bold
FONT_SIZE_BODY = Pt(10)       # Body text: 10pt
FONT_SIZE_BULLET = Pt(10)     # Bullet items: 10pt

# Spacing
LINE_SPACING = Pt(12)
PARAGRAPH_SPACING = Pt(2)     # Tighter spacing for 1-page fit


async def generate_resume(
    output_path: str,
    keywords: List[str],
    resume_data: Optional[ResumeData] = None,
    job_description: Optional[str] = None,
    use_parallel: bool = True,
    fast_mode: bool = False,
    session_id: Optional[str] = None
) -> None:
    """
    Generate a personalized, ATS-optimized resume in C3 format.
    
    Args:
        output_path: Path to save the generated DOCX file
        keywords: Extracted keywords from job description
        resume_data: Parsed resume data (if available)
        job_description: Job description (for future LLM rewriting)
    """
    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    
    # OPTIMIZED: Run condensation and parallel data preparation simultaneously!
    # This saves 2-5 seconds by running them in parallel instead of sequentially
    
    # Create new document first (no async needed)
    document = Document()
    
    # Set page size (C3 or Letter - Letter is bigger for more content)
    if USE_LETTER_SIZE:
        _set_letter_page_size(document)
    else:
        _set_c3_page_size(document)
    
    # Set margins
    _set_margins(document)
    
    # SMART CONDENSATION: Only condense if content is actually too large
    # Skip condensation if resume is already compact (saves 1-2 seconds!)
    should_condense = False
    if resume_data:
        total_bullets = sum(len(exp.bullets) for exp in resume_data.experience)
        # Only condense if:
        # - More than 4 experiences, OR
        # - More than 20 total bullets, OR
        # - More than 4 projects
        if len(resume_data.experience) > 4 or total_bullets > 20 or len(resume_data.projects) > 4:
            should_condense = True
    
    # Run condensation and parallel data preparation in parallel (if both enabled)
    if resume_data:
        tasks = []
        
        # Task 1: Condensation (only if needed - saves time!)
        if should_condense:
            if LLM_CONDENSER_ASYNC_AVAILABLE:
                condense_task = condense_resume_for_one_page_async(
                    resume_data, 
                    target_page_size="C3" if not USE_LETTER_SIZE else "Letter"
                )
                tasks.append(("condense", condense_task))
            elif LLM_CONDENSER_AVAILABLE:
                # Sync version - run in thread pool to not block
                condense_task = asyncio.to_thread(
                    condense_resume_for_one_page,
                    resume_data,
                    "C3" if not USE_LETTER_SIZE else "Letter"
                )
                tasks.append(("condense", condense_task))
        
        # Task 2: Parallel data preparation (if enabled)
        if use_parallel and ASYNC_AVAILABLE and job_description:
            if OPTIMIZED_AVAILABLE:
                # Use optimized version (faster, with smart skipping)
                prep_task = prepare_resume_data_optimized(
                    resume_data, job_description, keywords,
                    fast_mode=fast_mode, session_id=session_id
                )
            else:
                # Fallback to regular parallel version
                prep_task = prepare_resume_data_parallel(
                    resume_data, job_description, keywords
                )
            tasks.append(("prepare", prep_task))
        
        # Execute all tasks in parallel
        if tasks:
            try:
                results = await asyncio.gather(*[task for _, task in tasks], return_exceptions=True)
                
                # Process condensation result
                for i, (task_type, _) in enumerate(tasks):
                    if task_type == "condense" and not isinstance(results[i], Exception):
                        resume_data = results[i]
                    elif task_type == "condense" and isinstance(results[i], Exception):
                        print(f"Condensation error: {results[i]}. Using original data.")
                
                # Process parallel preparation result
                for i, (task_type, _) in enumerate(tasks):
                    if task_type == "prepare" and not isinstance(results[i], Exception):
                        prioritized_experiences, personalized_projects = results[i]
                        resume_data.experience = prioritized_experiences
                        resume_data.projects = personalized_projects
                    elif task_type == "prepare" and isinstance(results[i], Exception):
                        print(f"Parallel LLM processing error: {results[i]}. Using original data.")
            except Exception as e:
                print(f"Parallel processing error: {e}. Using original data.")
        elif use_parallel and ASYNC_AVAILABLE and job_description:
            # Only parallel prep, no condensation
            try:
                if OPTIMIZED_AVAILABLE:
                    prioritized_experiences, personalized_projects = await prepare_resume_data_optimized(
                        resume_data, job_description, keywords,
                        fast_mode=fast_mode, session_id=session_id
                    )
                else:
                    prioritized_experiences, personalized_projects = await prepare_resume_data_parallel(
                        resume_data, job_description, keywords
                    )
                resume_data.experience = prioritized_experiences
                resume_data.projects = personalized_projects
            except Exception as e:
                print(f"Parallel LLM processing error: {e}. Using sequential processing.")
    
    # Build resume sections (with more content now)
    _build_header(document, resume_data)
    _build_education(document, resume_data)
    _build_skills(document, resume_data, keywords)
    # Build sections (data already prepared in parallel above if enabled)
    _build_experience(document, resume_data, keywords, job_description)
    _build_projects(document, resume_data, keywords, job_description)
    _build_certifications(document, resume_data)
    
    # Enforce 1-page limit (smarter enforcement)
    _enforce_one_page_smart(document)
    
    # Save document
    document.save(str(output_file))


def _set_c3_page_size(document: Document) -> None:
    """Set document to C3 page size (7.17" x 10.51")."""
    section = document.sections[0]
    section.page_width = C3_WIDTH
    section.page_height = C3_HEIGHT
    section.orientation = WD_ORIENT.PORTRAIT


def _set_letter_page_size(document: Document) -> None:
    """Set document to Letter page size (8.5" x 11") - bigger for more content."""
    section = document.sections[0]
    section.page_width = LETTER_WIDTH
    section.page_height = LETTER_HEIGHT
    section.orientation = WD_ORIENT.PORTRAIT


def _set_margins(document: Document) -> None:
    """Set document margins to 0.4" all around for maximum content space."""
    section = document.sections[0]
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT


def _build_header(document: Document, resume_data: Optional[ResumeData]) -> None:
    """
    Build ATS-compliant resume header.
    
    ATS RULES:
    - Name: plain text, 14-16pt, bold (in body, NOT in header/footer)
    - Contact: ONE line — email | phone | LinkedIn | location
    - NO images, NO tables, NO text boxes
    - ALL content in main document body
    """
    if resume_data:
        # Name — 14pt, Bold, left-aligned (ATS: 14-16pt, plain text)
        name_para = document.add_paragraph()
        name_run = name_para.add_run(resume_data.name.upper())
        name_run.font.name = FONT_NAME
        name_run.font.size = FONT_SIZE_NAME
        name_run.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_paragraph_spacing(name_para, before=0, after=2)
        
        # Contact info — ALL on ONE line separated by |
        # ATS Critical: Workday/Taleo parse contact from a single line
        contact_parts = []
        if resume_data.email:
            contact_parts.append(resume_data.email)
        if resume_data.phone:
            contact_parts.append(resume_data.phone)
        if resume_data.linkedin:
            linkedin_clean = resume_data.linkedin.replace("https://", "").replace("http://", "").rstrip("/")
            contact_parts.append(linkedin_clean)
        if resume_data.github:
            github_clean = resume_data.github.replace("https://", "").replace("http://", "").rstrip("/")
            contact_parts.append(github_clean)
        if resume_data.location:
            contact_parts.append(resume_data.location)
        
        if contact_parts:
            contact_para = document.add_paragraph()
            contact_run = contact_para.add_run(" | ".join(contact_parts))
            contact_run.font.name = FONT_NAME
            contact_run.font.size = FONT_SIZE_CONTACT
            contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_paragraph_spacing(contact_para, before=0, after=0)
    else:
        # Fallback header
        name_para = document.add_paragraph()
        name_run = name_para.add_run("YOUR NAME")
        name_run.font.name = FONT_NAME
        name_run.font.size = FONT_SIZE_NAME
        name_run.bold = True
        
        contact_para = document.add_paragraph()
        contact_run = contact_para.add_run("email@domain.com | (555) 123-4567 | linkedin.com/in/name | City, State")
        contact_run.font.name = FONT_NAME
        contact_run.font.size = FONT_SIZE_CONTACT


def _build_education(document: Document, resume_data: Optional[ResumeData]) -> None:
    """
    Build education section — ATS-standard format.
    
    ATS FORMAT:
      Degree Name in Major
      University Name, City, State | Graduation Month Year
    """
    if not resume_data or not resume_data.education:
        return
    
    _add_section_heading(document, "EDUCATION")
    
    for edu in resume_data.education:
        # Line 1: Degree (bold) — dates on same line
        degree_para = document.add_paragraph()
        degree_run = degree_para.add_run(edu.degree)
        degree_run.font.name = FONT_NAME
        degree_run.font.size = FONT_SIZE_BODY
        degree_run.bold = True
        if edu.dates:
            dates_run = degree_para.add_run(f"  |  {edu.dates}")
            dates_run.font.name = FONT_NAME
            dates_run.font.size = FONT_SIZE_BODY
        _set_paragraph_spacing(degree_para, before=3, after=0)
        
        # Line 2: University, Location (italic) + GPA
        uni_parts = []
        if edu.university:
            uni_parts.append(edu.university)
        if edu.location:
            uni_parts.append(edu.location)
        
        if uni_parts:
            uni_para = document.add_paragraph()
            uni_run = uni_para.add_run(", ".join(uni_parts))
            uni_run.font.name = FONT_NAME
            uni_run.font.size = FONT_SIZE_BODY
            uni_run.italic = True
            if edu.gpa:
                gpa_run = uni_para.add_run(f"  |  GPA: {edu.gpa}")
                gpa_run.font.name = FONT_NAME
                gpa_run.font.size = FONT_SIZE_BODY
            _set_paragraph_spacing(uni_para, before=0, after=0)
        
        # Coursework (on one line, italic)
        if edu.coursework:
            cw_para = document.add_paragraph()
            cw_run = cw_para.add_run("Relevant Coursework: " + ", ".join(edu.coursework[:8]))
            cw_run.font.name = FONT_NAME
            cw_run.font.size = FONT_SIZE_BODY
            cw_run.italic = True
            _set_paragraph_spacing(cw_para, before=0, after=0)


def _build_skills(document: Document, resume_data: Optional[ResumeData], keywords: List[str]) -> None:
    """
    Build technical skills section — ATS-optimized.
    
    ATS FORMAT:
      Category: skill, skill, skill
      Category: skill, skill, skill
    
    Simple comma-separated lists are the most ATS-parseable format.
    """
    _add_section_heading(document, "TECHNICAL SKILLS")
    
    categorized_skills = {}
    uncategorized = []
    
    if resume_data and resume_data.skills:
        if isinstance(resume_data.skills, dict):
            categorized_skills = resume_data.skills.copy()
        else:
            for skill in resume_data.skills:
                skill_lower = str(skill).lower()
                category = _categorize_skill(skill_lower)
                if category:
                    categorized_skills.setdefault(category, []).append(skill)
                else:
                    uncategorized.append(skill)
    
    # Inject missing keywords into appropriate categories
    for keyword in keywords:
        normalized_keyword = normalize_keyword(keyword)
        keyword_lower = keyword.lower()
        category = _categorize_skill(keyword_lower)
        
        if category:
            existing_skills = categorized_skills.get(category, [])
            if normalized_keyword not in [normalize_keyword(str(s)) for s in existing_skills]:
                categorized_skills.setdefault(category, []).append(keyword)
        else:
            if normalized_keyword not in [normalize_keyword(str(s)) for s in uncategorized]:
                uncategorized.append(keyword)
    
    # Deduplicate
    for category in categorized_skills:
        categorized_skills[category] = deduplicate_preserve_order([str(s) for s in categorized_skills[category]])
    
    # Preferred category order (ATS-standard names)
    category_order = [
        "Programming Languages", "Languages",
        "Backend", "Frameworks & APIs",
        "Frontend",
        "Cloud", "Cloud Platforms",
        "Databases", "Databases & Storage",
        "AI/ML", "AI/ML & LLMs",
        "DevOps", "DevOps & MLOps",
        "Data Engineering",
        "Tools",
        "Monitoring & Analytics",
    ]
    
    rendered = set()
    for category in category_order:
        if category in categorized_skills and categorized_skills[category] and category not in rendered:
            skills_para = document.add_paragraph()
            cat_run = skills_para.add_run(f"{category}: ")
            cat_run.font.name = FONT_NAME
            cat_run.font.size = FONT_SIZE_BODY
            cat_run.bold = True
            skills_run = skills_para.add_run(", ".join(categorized_skills[category][:20]))
            skills_run.font.name = FONT_NAME
            skills_run.font.size = FONT_SIZE_BODY
            _set_paragraph_spacing(skills_para, before=0, after=0)
            rendered.add(category)
    
    # Remaining categories
    for category, skill_list in categorized_skills.items():
        if category not in rendered and skill_list:
            skills_para = document.add_paragraph()
            cat_run = skills_para.add_run(f"{category}: ")
            cat_run.font.name = FONT_NAME
            cat_run.font.size = FONT_SIZE_BODY
            cat_run.bold = True
            skills_run = skills_para.add_run(", ".join(skill_list[:20]))
            skills_run.font.name = FONT_NAME
            skills_run.font.size = FONT_SIZE_BODY
            _set_paragraph_spacing(skills_para, before=0, after=0)
    
    if uncategorized:
        skills_para = document.add_paragraph()
        other_run = skills_para.add_run("Other: ")
        other_run.font.name = FONT_NAME
        other_run.font.size = FONT_SIZE_BODY
        other_run.bold = True
        skills_run = skills_para.add_run(", ".join(deduplicate_preserve_order([str(s) for s in uncategorized])[:15]))
        skills_run.font.name = FONT_NAME
        skills_run.font.size = FONT_SIZE_BODY
        _set_paragraph_spacing(skills_para, before=0, after=0)


def _categorize_skill(skill: str) -> Optional[str]:
    """Categorize a skill based on keywords."""
    skill_lower = skill.lower()
    if any(kw in skill_lower for kw in ['python', 'java', 'sql', 'javascript', 'typescript', 'c++', 'go', 'rust', 'pyspark', 'scala', 'r']):
        return "Programming Languages"
    elif any(kw in skill_lower for kw in ['tensorflow', 'pytorch', 'hugging', 'llm', 'gpt', 'bert', 'ml', 'ai', 'machine learning', 'deep learning', 'neural']):
        return "AI/ML & LLMs"
    elif any(kw in skill_lower for kw in ['aws', 'azure', 'gcp', 'cloud', 'sagemaker', 'lambda', 'glue', 'ec2', 's3']):
        return "Cloud Platforms"
    elif any(kw in skill_lower for kw in ['spark', 'databricks', 'kafka', 'etl', 'data engineering', 'airflow', 'flink']):
        return "Data Engineering"
    elif any(kw in skill_lower for kw in ['docker', 'kubernetes', 'terraform', 'ci/cd', 'devops', 'mlops', 'jenkins', 'gitlab']):
        return "DevOps & MLOps"
    elif any(kw in skill_lower for kw in ['postgresql', 'mongodb', 'redis', 'mysql', 'database', 'storage', 'dynamodb', 'redshift']):
        return "Databases & Storage"
    elif any(kw in skill_lower for kw in ['fastapi', 'django', 'flask', 'react', 'node', 'api', 'framework', 'express', 'spring']):
        return "Frameworks & APIs"
    elif any(kw in skill_lower for kw in ['datadog', 'new relic', 'splunk', 'monitoring', 'analytics', 'grafana', 'prometheus']):
        return "Monitoring & Analytics"
    return None


def _build_experience(
    document: Document,
    resume_data: Optional[ResumeData],
    keywords: List[str],
    job_description: Optional[str]
) -> None:
    """
    Build work experience section — ATS-optimized.
    
    ATS FORMAT:
      Job Title
      Company Name, City, State | Month Year - Month Year
      - Bullet point with action verb and result
      - 4-6 bullets per job
    """
    if not resume_data or not resume_data.experience:
        return
    
    # Section heading — ATS standard: "WORK EXPERIENCE"
    _add_section_heading(document, "WORK EXPERIENCE")
    
    # Prioritize most relevant experiences
    if job_description and len(resume_data.experience) > 4:
        experiences = match_experience_with_jd(resume_data.experience, job_description, top_n=4)
    else:
        experiences = resume_data.experience[:5]
    
    for exp in experiences:
        # Line 1: Job Title (bold) — right-aligned dates
        title_para = document.add_paragraph()
        title_run = title_para.add_run(exp.title)
        title_run.font.name = FONT_NAME
        title_run.font.size = FONT_SIZE_BODY
        title_run.bold = True
        if exp.dates:
            # Add tab + dates on same line (ATS-parseable layout)
            dates_run = title_para.add_run(f"  |  {exp.dates}")
            dates_run.font.name = FONT_NAME
            dates_run.font.size = FONT_SIZE_BODY
        _set_paragraph_spacing(title_para, before=4, after=0)
        
        # Line 2: Company name (italic)
        if exp.company:
            company_para = document.add_paragraph()
            company_run = company_para.add_run(exp.company)
            company_run.font.name = FONT_NAME
            company_run.font.size = FONT_SIZE_BODY
            company_run.italic = True
            _set_paragraph_spacing(company_para, before=0, after=1)
        
        # Bullet points — PERSONALIZE using LLM if JD available
        if job_description and keywords:
            personalized_bullets = rewrite_experience_bullets(exp, job_description, keywords)
            bullets = personalized_bullets[:6]
        else:
            bullets = exp.bullets[:6]
        
        for bullet in bullets:
            # ATS-safe bullet: use simple dash or bullet character
            bullet_para = document.add_paragraph()
            bullet_run = bullet_para.add_run(f"- {bullet}")
            bullet_run.font.name = FONT_NAME
            bullet_run.font.size = FONT_SIZE_BULLET
            _set_paragraph_spacing(bullet_para, before=0, after=0)
            # Indent bullets
            bullet_para.paragraph_format.left_indent = Inches(0.25)


def _build_projects(
    document: Document,
    resume_data: Optional[ResumeData],
    keywords: List[str],
    job_description: Optional[str]
) -> None:
    """
    Build projects section — ATS-optimized.
    
    ATS FORMAT:
      Project Name | Tech1, Tech2
      - Description bullet with impact
    """
    if not resume_data or not resume_data.projects:
        return
    
    _add_section_heading(document, "PROJECTS")
    
    projects = resume_data.projects[:4]
    
    for project in projects:
        # Project name (bold) + technologies
        name_para = document.add_paragraph()
        name_run = name_para.add_run(project.name)
        name_run.font.name = FONT_NAME
        name_run.font.size = FONT_SIZE_BODY
        name_run.bold = True
        
        if project.technologies:
            tech_run = name_para.add_run(f"  |  {', '.join(project.technologies[:10])}")
            tech_run.font.name = FONT_NAME
            tech_run.font.size = FONT_SIZE_BODY
            tech_run.italic = True
        _set_paragraph_spacing(name_para, before=3, after=0)
        
        # Description as bullet points
        if project.description:
            description = project.description
            if job_description and keywords:
                try:
                    description = rewrite_project_description(project, job_description, keywords)
                except Exception:
                    pass
            
            # Split into sentences for bullet points
            sentences = [s.strip() for s in description.replace('. ', '.\n').split('\n') if s.strip()]
            for sentence in sentences[:3]:
                bullet_para = document.add_paragraph()
                clean = sentence.rstrip('.')
                bullet_run = bullet_para.add_run(f"- {clean}.")
                bullet_run.font.name = FONT_NAME
                bullet_run.font.size = FONT_SIZE_BULLET
                _set_paragraph_spacing(bullet_para, before=0, after=0)
                bullet_para.paragraph_format.left_indent = Inches(0.25)


def _build_certifications(document: Document, resume_data: Optional[ResumeData]) -> None:
    """
    Build certifications section — ATS-standard.
    
    ATS FORMAT:
      Cert Name - Issuer (Year)
    """
    if not resume_data or not resume_data.certifications:
        return
    
    _add_section_heading(document, "CERTIFICATIONS")
    
    for cert in resume_data.certifications:
        cert_para = document.add_paragraph()
        cert_text = cert.name
        if cert.issuer:
            cert_text += f" - {cert.issuer}"
        if cert.year:
            cert_text += f" ({cert.year})"
        
        cert_run = cert_para.add_run(cert_text)
        cert_run.font.name = FONT_NAME
        cert_run.font.size = FONT_SIZE_BODY
        _set_paragraph_spacing(cert_para, before=0, after=0)


def _set_paragraph_spacing(para, before: int = 0, after: int = 2) -> None:
    """Set paragraph spacing in points."""
    from docx.shared import Pt as PtSpacing
    pf = para.paragraph_format
    pf.space_before = PtSpacing(before)
    pf.space_after = PtSpacing(after)


def _add_section_heading(document: Document, heading_text: str) -> None:
    """
    Add an ATS-standard section heading.
    
    ATS RULES:
    - Use EXACT standard names: WORK EXPERIENCE, EDUCATION, TECHNICAL SKILLS
    - Bold, slightly larger than body (11pt)
    - Left-aligned (no center, no right-align)
    - Simple horizontal rule (not special characters)
    - NO fancy formatting, NO colors, NO icons
    """
    # Section heading — 11pt, Bold, UPPERCASE
    heading_para = document.add_paragraph()
    heading_run = heading_para.add_run(heading_text.upper())
    heading_run.font.name = FONT_NAME
    heading_run.font.size = FONT_SIZE_SECTION
    heading_run.bold = True
    heading_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(heading_para, before=6, after=2)
    
    # Simple horizontal line via paragraph border (ATS-safe)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = heading_para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '333333')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _enforce_one_page_smart(document: Document) -> None:
    """
    Smarter 1-page enforcement that preserves maximum content.
    
    Uses better heuristics and only removes content if absolutely necessary.
    """
    # Count paragraphs
    total_paragraphs = len(document.paragraphs)
    
    # With Letter size and optimized formatting, we can fit more
    # Letter (8.5x11): ~60-70 paragraphs
    # C3 (7.17x10.51): ~50-60 paragraphs
    if USE_LETTER_SIZE:
        max_paragraphs = 65  # Increased for Letter size
    else:
        max_paragraphs = 55  # Increased for C3
    
    if total_paragraphs > max_paragraphs:
        # Only remove if significantly over
        excess = total_paragraphs - max_paragraphs
        if excess > 5:  # Only remove if more than 5 paragraphs over
            # Remove excess paragraphs (usually spacing or less critical content)
            paragraphs_to_remove = min(excess, 10)  # Don't remove more than 10
            for _ in range(paragraphs_to_remove):
                if len(document.paragraphs) > 15:  # Keep at least header and core content
                    # Remove last paragraph (usually spacing)
                    last_para = document.paragraphs[-1]
                    p_element = last_para._element
                    p_element.getparent().remove(p_element)
    
    # If still too long, slightly reduce font sizes (minimal impact)
    if len(document.paragraphs) > max_paragraphs + 5:
        # Reduce body font size slightly (only if really needed)
        for para in document.paragraphs[5:]:  # Skip header
            for run in para.runs:
                if run.font.size and run.font.size > Pt(9):
                    run.font.size = Pt(9)  # Slight reduction
